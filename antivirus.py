#!/usr/bin/env python3
"""
Odyssey Antivirus
National Teachers College - Information Assurance and Security 1 Finals Project

Advanced antivirus tool specifically designed to detect and remove
the Odyssey Virus and demonstrate cybersecurity principles.

Author: Tavera'S Group  
Institution: National Teachers College
Date: June 2025
Target: Odyssey Virus

MODIFICATIONS:
- Focus on encrypted files detection and recovery
- Handle deleted original files scenario
- Enhanced recovery capabilities for replaced files
- Optimized scanning for encrypted content only

FEATURES:
- Multi-format document processing
- Encrypted file detection and analysis
- Advanced behavioral analysis targeting encrypted files
- Document-aware quarantine system for encrypted files
- ROT13 decryption engine
- Comprehensive virus signature detection for encrypted content
- Secure file recovery and restoration
- Real-time scanning with GUI
- Detailed forensic reporting
"""

import os
import sys
import json
import time
import hashlib
import fnmatch
import threading
from datetime import datetime
from pathlib import Path

# Add explicit debugging
print("DEBUG: Starting antivirus import process...")

try:
    import tkinter as tk
    from tkinter import messagebox, scrolledtext, filedialog, ttk
    GUI_AVAILABLE = True
    print("DEBUG: tkinter imported successfully")
except ImportError as e:
    GUI_AVAILABLE = False
    print(f"DEBUG: tkinter import failed: {e}")

# Document processing libraries
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
    print("DEBUG: python-docx imported successfully")
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"DEBUG: python-docx not available: {e}")

try:
    import PyPDF2
    PDF_AVAILABLE = True
    print("DEBUG: PyPDF2 imported successfully")
except ImportError as e:
    PDF_AVAILABLE = False
    print(f"DEBUG: PyPDF2 not available: {e}")

try:
    import openpyxl
    XLSX_AVAILABLE = True
    print("DEBUG: openpyxl imported successfully")
except ImportError as e:
    XLSX_AVAILABLE = False
    print(f"DEBUG: openpyxl not available: {e}")

try:
    from pptx import Presentation
    PPTX_AVAILABLE = True
    print("DEBUG: python-pptx imported successfully")
except ImportError as e:
    PPTX_AVAILABLE = False
    print(f"DEBUG: python-pptx not available: {e}")

print("DEBUG: All imports completed")

class ROT13CryptographyEngine:
    """
    ROT13 Cryptographic Engine for Virus Decryption
    
    Implements the exact ROT13 algorithm used by the Odyssey virus
    to ensure proper decryption and file recovery.
    """
    
    @staticmethod
    def rot13_transform(text):
        """Apply ROT13 transformation (identical to virus implementation)"""
        result = []
        
        for char in text:
            if char.isalpha():
                is_upper = char.isupper()
                char = char.upper()
                transformed_char = chr(((ord(char) - ord('A') + 13) % 26) + ord('A'))
                
                if not is_upper:
                    transformed_char = transformed_char.lower()
                    
                result.append(transformed_char)
            else:
                result.append(char)
        
        return ''.join(result)
    
    @staticmethod
    def rot13_decrypt(ciphertext):
        """Decrypt ROT13 ciphertext (ROT13 is self-inverse)"""
        return ROT13CryptographyEngine.rot13_transform(ciphertext)
    
    @staticmethod
    def rot13_encrypt(plaintext):
        """Encrypt plaintext using ROT13 (for testing purposes)"""
        return ROT13CryptographyEngine.rot13_transform(plaintext)

class DocumentAnalyzer:
    """
    Advanced document analysis engine for encrypted files
    
    Provides analysis and recovery capabilities for:
    - Encrypted Word Documents (.docx)
    - Encrypted PDF files (.pdf)
    - Encrypted Excel files (.xlsx)
    - Encrypted PowerPoint presentations (.pptx)
    - Encrypted plain text files
    """
    
    def __init__(self, crypto_engine):
        self.crypto_engine = crypto_engine
        self.supported_formats = {
            '.docx': 'Word Document',
            '.pdf': 'PDF Document', 
            '.xlsx': 'Excel Spreadsheet',
            '.pptx': 'PowerPoint Presentation',
            '.txt': 'Text File',
            '.md': 'Markdown File',
            '.csv': 'CSV File',
            '.py': 'Python Script',
            '.js': 'JavaScript File',
            '.html': 'HTML File',
            '.css': 'CSS File'
        }
    
    def analyze_encrypted_word_document(self, file_path):
        """Analyze encrypted Word document for virus content"""
        if not DOCX_AVAILABLE:
            return None, "python-docx library not available"
        
        try:
            doc = Document(file_path)
            text_content = []
            virus_indicators = []
            
            # Extract and analyze paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if text.strip():
                    text_content.append(text)
                    
                    # Check for virus signatures
                    if "ODYSSEY_VIRUS_2025_NTC" in text:
                        virus_indicators.append(f"Virus signature found: {text[:100]}...")
                    if "ROT13" in text and "Encryption" in text:
                        virus_indicators.append(f"Encryption indicator: {text[:100]}...")
            
            # Analyze tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip() and "ODYSSEY" in cell.text:
                            virus_indicators.append(f"Table virus content: {cell.text[:50]}...")
            
            analysis_result = {
                'document_type': 'Encrypted Word Document',
                'paragraphs_count': len([p for p in doc.paragraphs if p.text.strip()]),
                'tables_count': len(doc.tables),
                'text_content': '\n'.join(text_content),
                'virus_indicators': virus_indicators,
                'is_infected': len(virus_indicators) > 0,
                'encrypted': True
            }
            
            return analysis_result, None
            
        except Exception as e:
            return None, f"Error analyzing encrypted Word document: {str(e)}"
    
    def analyze_encrypted_text_file(self, file_path):
        """Analyze encrypted text file content"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            virus_indicators = []
            if "ODYSSEY_VIRUS_2025_NTC" in content:
                virus_indicators.append("Virus signature detected")
            if "ROT13" in content and "encrypted" in content.lower():
                virus_indicators.append("Encryption indicators found")
            if "=== ENCRYPTED" in content:
                virus_indicators.append("Encrypted content marker found")
            if "ODYSSEY_ENCRYPTED" in file_path:
                virus_indicators.append("Encrypted filename pattern detected")
            
            analysis_result = {
                'document_type': 'Encrypted Text File',
                'text_content': content,
                'virus_indicators': virus_indicators,
                'is_infected': len(virus_indicators) > 0,
                'encrypted': True
            }
            
            return analysis_result, None
            
        except Exception as e:
            return None, f"Error reading encrypted text file: {str(e)}"
    
    def analyze_document(self, file_path):
        """Analyze document based on file extension and encryption status"""
        filename = os.path.basename(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Prioritize encrypted files
        if "ODYSSEY_ENCRYPTED" in filename:
            if file_extension == '.docx':
                return self.analyze_encrypted_word_document(file_path)
            else:
                return self.analyze_encrypted_text_file(file_path)
        
        # Analyze non-encrypted files (for completeness)
        if file_extension == '.docx' and DOCX_AVAILABLE:
            try:
                doc = Document(file_path)
                text_content = []
                virus_indicators = []
                
                for paragraph in doc.paragraphs:
                    text = paragraph.text
                    if text.strip():
                        text_content.append(text)
                        if "ODYSSEY_VIRUS_2025_NTC" in text:
                            virus_indicators.append(f"Virus signature found: {text[:100]}...")
                
                analysis_result = {
                    'document_type': 'Word Document',
                    'paragraphs_count': len([p for p in doc.paragraphs if p.text.strip()]),
                    'text_content': '\n'.join(text_content),
                    'virus_indicators': virus_indicators,
                    'is_infected': len(virus_indicators) > 0,
                    'encrypted': False
                }
                
                return analysis_result, None
                
            except Exception as e:
                return None, f"Error analyzing Word document: {str(e)}"
        
        elif file_extension in ['.txt', '.md', '.csv', '.py', '.js', '.html', '.css']:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                
                virus_indicators = []
                if "ODYSSEY_VIRUS_2025_NTC" in content:
                    virus_indicators.append("Virus signature detected")
                
                analysis_result = {
                    'document_type': self.supported_formats.get(file_extension, 'Text File'),
                    'text_content': content,
                    'virus_indicators': virus_indicators,
                    'is_infected': len(virus_indicators) > 0,
                    'encrypted': False
                }
                
                return analysis_result, None
                
            except Exception as e:
                return None, f"Error reading text file: {str(e)}"
        else:
            return None, f"Unsupported file type: {file_extension}"

class OdysseyVirusSignatureDatabase:
    """
    Signature database for multi-document Odyssey virus detection
    """
    
    def __init__(self):
        self.virus_signatures = {
            "ODYSSEY_VIRUS_2025_NTC": {
                "name": "Odyssey Virus",
                "version": "2.1_EDUCATIONAL",
                "institution": "National Teachers College",
                "type": "Multi-Document Educational Malware with File Replacement",
                "risk_level": "Educational Only",
                "encryption_algorithm": "ROT13",
                "target_directory": "odyssey_test_environment",
                "supported_formats": [".docx", ".pdf", ".xlsx", ".pptx", ".txt", ".md", ".csv"],
                "description": "ROT13-based virus with multi-document support and file replacement behavior",
                "behavioral_characteristics": ["File replacement", "Original file deletion", "Encrypted file creation"]
            }
        }
        
        self.encrypted_file_patterns = [
            "*ODYSSEY_ENCRYPTED*",           # All encrypted files
            "*ODYSSEY_LOCKED_*",             # Temporarily locked files
        ]
        
        self.infected_file_indicators = [
            "ODYSSEY_VIRUS_2025_NTC INFECTION MARKER",
            "Educational Malware Injection",
            "Virus Code Injection",
            "Script Infection",
            "INFECTED BY ODYSSEY_VIRUS_2025_NTC"
        ]
        
        self.artifact_patterns = [
            "odyssey_activity.log",          # Activity log
            ".odyssey_infection_marker",     # Infection marker
            "odyssey_encrypted_payload.dat", # Encrypted payload
            "odyssey_encryption_manifest.json", # Encryption manifest
            "odyssey_infection_manifest.json", # Infection manifest
            "*.backup",                      # Backup files created during infection
        ]
        
        self.content_signatures = [
            "ODYSSEY_VIRUS_2025_NTC",
            "<!-- ODYSSEY_VIRUS_2025_NTC -->",
            "<!-- Encryption: ROT13 -->",
            "National Teachers College",
            "Odyssey Virus",
            "ROT13 Cryptographic Implementation",
            "=== ENCRYPTED .DOCX CONTENT ===",
            "=== ENCRYPTED .PDF CONTENT ===",
            "=== ENCRYPTED .XLSX CONTENT ===",
            "=== ENCRYPTED .PPTX CONTENT ===",
            "=== ENCRYPTED",
            "Original File:",
            "Multi-Document Support",
            "File replacement simulation"
        ]
        
        self.encrypted_document_indicators = {
            'encrypted_word_document': [
                "Encrypted Document - ROT13",
                "Educational Note:",
                "This document has been encrypted using ROT13"
            ],
            'encrypted_generic': [
                "=== ENCRYPTED",
                "Original File:",
                "Encryption: ROT13",
                "Document Type:"
            ]
        }
        
        self.behavioral_indicators = [
            "File encryption with ODYSSEY_ENCRYPTED suffix",
            "ROT13 encrypted educational messages",
            "Infection marker placement",
            "Temporary file locking with ODYSSEY_LOCKED prefix",
            "Encrypted payload generation",
            "Educational popup message display",
            "Multi-format document processing",
            "Original file deletion and replacement",
            "Comprehensive activity logging"
        ]
    
    def get_virus_info(self, signature):
        """Get detailed information about virus signature"""
        return self.virus_signatures.get(signature, None)
    
    def is_encrypted_file_pattern(self, filename):
        """Check if filename matches encrypted file patterns"""
        for pattern in self.encrypted_file_patterns:
            if fnmatch.fnmatch(filename, pattern):
                return True
        return False
    
    def is_artifact_pattern(self, filename):
        """Check if filename matches virus artifact patterns"""
        for pattern in self.artifact_patterns:
            if fnmatch.fnmatch(filename, pattern):
                return True
        return False
    
    def is_file_infected(self, content):
        """Check if file content contains infection indicators"""
        for indicator in self.infected_file_indicators:
            if indicator in content:
                return True
        return False
    
    def get_infection_indicators(self, content):
        """Get list of infection indicators found in content"""
        indicators = []
        for indicator in self.infected_file_indicators:
            if indicator in content:
                indicators.append(indicator)
        return indicators
    
    def analyze_encrypted_content(self, content):
        """Analyze encrypted content for infection indicators"""
        indicators = []
        
        for signature in self.content_signatures:
            if signature in content:
                indicators.append(signature)
        
        # Check for document-specific indicators
        for doc_type, patterns in self.encrypted_document_indicators.items():
            for pattern in patterns:
                if pattern in content:
                    indicators.append(f"{doc_type}: {pattern}")
        
        return indicators
    
    def get_all_signatures(self):
        """Get all known virus signatures"""
        return list(self.virus_signatures.keys())

class AdvancedVirusScanner:
    """
    Virus scanning engine with focus on encrypted files
    
    Implements multiple detection methods optimized for encrypted content:
    - Encrypted file signature detection
    - ROT13 content analysis
    - Behavioral pattern recognition for replaced files
    - Document-specific encrypted analysis
    - File integrity verification
    """
    
    def __init__(self):
        self.crypto_engine = ROT13CryptographyEngine()
        self.document_analyzer = DocumentAnalyzer(self.crypto_engine)
        self.signature_db = OdysseyVirusSignatureDatabase()
        self.scan_results = []
        self.quarantine_dir = "odyssey_antivirus_quarantine"
        self.scan_statistics = {
            'files_scanned': 0,
            'encrypted_files_found': 0,
            'infected_files_found': 0,
            'artifacts_found': 0,
            'threats_detected': 0,
            'document_types_scanned': {},
            'files_quarantined': 0,
            'files_cleaned': 0,
            'scan_duration': 0
        }
        
    def initialize_quarantine_system(self):
        """Initialize secure quarantine directory"""
        if not os.path.exists(self.quarantine_dir):
            os.makedirs(self.quarantine_dir)
            # Create quarantine metadata file
            quarantine_info = {
                "created": datetime.now().isoformat(),
                "purpose": "Odyssey Virus Quarantine",
                "antivirus": "Odyssey Virus Hunter",
                "institution": "National Teachers College",
                "supported_formats": list(self.document_analyzer.supported_formats.keys()),
                "target_focus": "Encrypted files and virus artifacts"
            }
            
            info_path = os.path.join(self.quarantine_dir, "quarantine_info.json")
            with open(info_path, 'w', encoding='utf-8') as f:
                json.dump(quarantine_info, f, indent=2)

            self.log_activity(f"Odyssey quarantine system initialized: {self.quarantine_dir}")
        else:
            self.log_activity(f"Using existing Odyssey quarantine: {self.quarantine_dir}")

    def log_activity(self, message, level="INFO"):
        """Logging with different levels"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Simple text prefixes for Windows compatibility
        level_prefixes = {
            "INFO": "[INFO]",
            "WARNING": "[WARN]",
            "ERROR": "[ERROR]",
            "SUCCESS": "[SUCCESS]",
            "DETECTION": "[DETECT]",
            "CRYPTO": "[CRYPTO]",
            "REMOVAL": "[REMOVE]",
            "DOCUMENT": "[DOC]",
            "ENCRYPTED": "[ENCRYPT]"
        }
        
        prefix = level_prefixes.get(level, "[INFO]")
        formatted_message = f"[{timestamp}] {prefix} {message}"
        print(formatted_message)
        
        # Also log to file if needed
        try:
            log_file = "antivirus_activity.log"
            with open(log_file, 'a', encoding='utf-8') as f:
                f.write(f"{formatted_message}\n")
        except:
            pass  # Don't let logging failures break the antivirus
    
    def calculate_file_hash(self, filepath):
        """Calculate SHA-256 hash of file for integrity verification"""
        try:
            hash_sha256 = hashlib.sha256()
            with open(filepath, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            self.log_activity(f"Failed to calculate hash for {filepath}: {str(e)}", "ERROR")
            return None
    
    def scan_encrypted_file_comprehensive(self, filepath):
        """Comprehensive encrypted file scanning"""
        detections = []
        filename = os.path.basename(filepath)
        file_extension = os.path.splitext(filepath)[1].lower()
        
        try:
            # Skip files outside odyssey_test_environment
            if "odyssey_test_environment" not in os.path.abspath(filepath):
                return detections
                
            # Skip non-files and quarantine directory
            if not os.path.isfile(filepath) or self.quarantine_dir in filepath:
                return detections
            
            # Update statistics
            self.scan_statistics['files_scanned'] += 1
            
            # Calculate file hash
            file_hash = self.calculate_file_hash(filepath)
            
            # Prioritize encrypted files
            if "ODYSSEY_ENCRYPTED" in filename:
                self.scan_statistics['encrypted_files_found'] += 1
                self.log_activity(f"Found encrypted file: {filename}", "ENCRYPTED")
                
                # Document-specific analysis for encrypted files
                if file_extension in self.document_analyzer.supported_formats:
                    analysis_result, error = self.document_analyzer.analyze_document(filepath)
                    
                    if error:
                        self.log_activity(f"Encrypted document analysis error for {filename}: {error}", "WARNING")
                    elif analysis_result and analysis_result.get('is_infected', False):
                        detection = {
                            'type': 'encrypted_document_infection',
                            'file': filepath,
                            'document_type': analysis_result['document_type'],
                            'virus_indicators': analysis_result['virus_indicators'],
                            'analysis_data': analysis_result,
                            'file_hash': file_hash,
                            'detection_time': datetime.now().isoformat(),
                            'confidence': 'HIGH',
                            'encrypted': True
                        }
                        detections.append(detection)
                        self.log_activity(f"Infected encrypted document: {filename} ({analysis_result['document_type']})", "DETECTION")
                
                # Always flag encrypted files as threats
                detection = {
                    'type': 'encrypted_file',
                    'file': filepath,
                    'file_hash': file_hash,
                    'detection_time': datetime.now().isoformat(),
                    'confidence': 'HIGH',
                    'description': f'Odyssey encrypted file: {filename}'
                }
                detections.append(detection)
                self.log_activity(f"Encrypted file detected: {filename}", "DETECTION")
            
            # Check for virus artifacts
            if self.signature_db.is_artifact_pattern(filename):
                self.scan_statistics['artifacts_found'] += 1
                detection = {
                    'type': 'virus_artifact',
                    'file': filepath,
                    'file_hash': file_hash,
                    'detection_time': datetime.now().isoformat(),
                    'confidence': 'HIGH',
                    'description': f'Odyssey virus artifact: {filename}'
                }
                detections.append(detection)
                self.log_activity(f"Virus artifact detected: {filename}", "DETECTION")
            
            # Content analysis for all text-readable files
            try:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                
                # Check for file infection
                if self.signature_db.is_file_infected(content):
                    self.scan_statistics['infected_files_found'] += 1
                    infection_indicators = self.signature_db.get_infection_indicators(content)
                    
                    detection = {
                        'type': 'infected_file',
                        'file': filepath,
                        'infection_indicators': infection_indicators,
                        'file_hash': file_hash,
                        'detection_time': datetime.now().isoformat(),
                        'confidence': 'HIGH',
                        'description': f'File infected with {len(infection_indicators)} virus signatures'
                    }
                    detections.append(detection)
                    self.log_activity(f"Infected file detected: {filename} ({len(infection_indicators)} indicators)", "DETECTION")
                
                # Check for virus signatures
                for signature in self.signature_db.get_all_signatures():
                    if signature in content:
                        virus_info = self.signature_db.get_virus_info(signature)
                        detection = {
                            'type': 'signature_match',
                            'signature': signature,
                            'file': filepath,
                            'file_hash': file_hash,
                            'virus_info': virus_info,
                            'detection_time': datetime.now().isoformat(),
                            'confidence': 'HIGH'
                        }
                        detections.append(detection)
                        self.log_activity(f"Virus signature detected: {signature} in {filename}", "DETECTION")
                
                # Analyze encrypted content indicators
                encryption_indicators = self.signature_db.analyze_encrypted_content(content)
                if encryption_indicators:
                    detection = {
                        'type': 'encrypted_content',
                        'file': filepath,
                        'indicators': encryption_indicators,
                        'file_hash': file_hash,
                        'detection_time': datetime.now().isoformat(),
                        'confidence': 'MEDIUM'
                    }
                    detections.append(detection)
                    self.log_activity(f"Encrypted content indicators: {len(encryption_indicators)} found in {filename}", "DETECTION")
                
                # ROT13 encryption detection
                if self.is_rot13_encrypted(content):
                    detection = {
                        'type': 'rot13_encrypted',
                        'file': filepath,
                        'file_hash': file_hash,
                        'detection_time': datetime.now().isoformat(),
                        'confidence': 'MEDIUM',
                        'description': 'File appears to contain ROT13 encrypted content'
                    }
                    detections.append(detection)
                    self.log_activity(f"ROT13 encryption detected in {filename}", "CRYPTO")
                    
            except Exception as e:
                # File might be binary or inaccessible
                pass
                
        except Exception as e:
            self.log_activity(f"Error scanning {filepath}: {str(e)}", "ERROR")
        
        return detections
    
    def is_rot13_encrypted(self, content):
        """Heuristic analysis to detect ROT13 encrypted content"""
        if len(content) < 50:  # Too short for reliable analysis
            return False
    
    def scan_filename_patterns(self, directory):
        """Scan directory for suspicious filename patterns (focus on encrypted files)"""
        detections = []
        
        try:
            for root, dirs, files in os.walk(directory):
                # Skip quarantine directory
                if self.quarantine_dir in root:
                    continue
                    
                for filename in files:
                    filepath = os.path.join(root, filename)
                    
                    # Check for encrypted file patterns
                    if self.signature_db.is_encrypted_file_pattern(filename):
                        detection = {
                            'type': 'encrypted_filename_pattern',
                            'pattern': filename,
                            'file': filepath,
                            'detection_time': datetime.now().isoformat(),
                            'confidence': 'HIGH',
                            'description': f'Filename matches Odyssey encrypted file pattern'
                        }
                        detections.append(detection)
                        self.log_activity(f"Encrypted filename pattern: {filename}", "DETECTION")
                    
                    # Check for artifact patterns
                    elif self.signature_db.is_artifact_pattern(filename):
                        detection = {
                            'type': 'artifact_filename_pattern',
                            'pattern': filename,
                            'file': filepath,
                            'detection_time': datetime.now().isoformat(),
                            'confidence': 'HIGH',
                            'description': f'Filename matches Odyssey artifact pattern'
                        }
                        detections.append(detection)
                        self.log_activity(f"Artifact filename pattern: {filename}", "DETECTION")
                        
        except Exception as e:
            self.log_activity(f"Error scanning filename patterns in {directory}: {str(e)}", "ERROR")
        
        return detections
    
    def analyze_infection_marker(self, directory):
        """Analyze Odyssey virus infection marker"""
        marker_file = os.path.join(directory, ".odyssey_infection_marker")
        
        if not os.path.exists(marker_file):
            return None
        
        try:
            with open(marker_file, 'r', encoding='utf-8') as f:
                marker_data = json.load(f)

            self.log_activity("Odyssey infection marker found - analyzing...", "DETECTION")
            
            # Extract and display key information
            if 'infection_metadata' in marker_data:
                metadata = marker_data['infection_metadata']
                self.log_activity(f"   Virus Signature: {metadata.get('virus_signature', 'Unknown')}")
                self.log_activity(f"   Virus Version: {metadata.get('virus_version', 'Unknown')}")
                self.log_activity(f"   Infection Time: {metadata.get('infection_timestamp', 'Unknown')}")
                
                if 'behavioral_modification' in metadata:
                    self.log_activity(f"   Behavior: {metadata.get('behavioral_modification', 'Unknown')}")
            
            if 'execution_statistics' in marker_data:
                stats = marker_data['execution_statistics']
                self.log_activity(f"   Files Processed: {stats.get('files_processed', 'Unknown')}")
                self.log_activity(f"   Files Replaced: {stats.get('files_replaced', 'Unknown')}")
                self.log_activity(f"   Encryption Operations: {stats.get('encryption_operations', 'Unknown')}")
                
                if 'document_types_processed' in stats:
                    doc_types = stats['document_types_processed']
                    self.log_activity(f"   Document Types Infected: {list(doc_types.keys())}")
            
            if 'behavioral_characteristics' in marker_data:
                behavior = marker_data['behavioral_characteristics']
                self.log_activity(f"   Original Files Deleted: {behavior.get('original_file_deletion', 'Unknown')}")
                self.log_activity(f"   File Replacement: {behavior.get('file_replacement', 'Unknown')}")
            
            return marker_data
            
        except Exception as e:
            self.log_activity(f"Failed to analyze infection marker: {str(e)}", "ERROR")
            return None
    
    def analyze_encrypted_payload(self, directory):
        """Analyze and decrypt Odyssey virus payload"""
        payload_file = os.path.join(directory, "odyssey_encrypted_payload.dat")
        
        if not os.path.exists(payload_file):
            return None
        
        try:
            with open(payload_file, 'r', encoding='utf-8') as f:
                encrypted_content = f.read()

            self.log_activity("Analyzing Odyssey encrypted payload...", "CRYPTO")
            
            # Remove virus signature comments if present
            lines = encrypted_content.split('\n')
            payload_content = []
            
            for line in lines:
                if not line.strip().startswith('<!--'):
                    payload_content.append(line)
            
            encrypted_payload = '\n'.join(payload_content).strip()
            
            # Decrypt using ROT13
            decrypted_payload = self.crypto_engine.rot13_decrypt(encrypted_payload)
            
            # Try to parse as JSON
            payload_data = json.loads(decrypted_payload)

            self.log_activity("Successfully decrypted Odyssey virus payload:", "SUCCESS")

            # Display key payload information
            if 'virus_identification' in payload_data:
                virus_id = payload_data['virus_identification']
                self.log_activity(f"   Virus Name: {virus_id.get('name', 'Unknown')}")
                self.log_activity(f"   Version: {virus_id.get('version', 'Unknown')}")
                self.log_activity(f"   Institution: {virus_id.get('institution', 'Unknown')}")
            
            if 'behavioral_changes' in payload_data:
                behavior = payload_data['behavioral_changes']
                self.log_activity("   Behavioral Changes:")
                self.log_activity(f"     File Deletion: {behavior.get('original_file_deletion', False)}")
                self.log_activity(f"     File Replacement: {behavior.get('file_replacement', False)}")
                self.log_activity(f"     Realistic Simulation: {behavior.get('realistic_simulation', False)}")
            
            if 'execution_metadata' in payload_data:
                exec_data = payload_data['execution_metadata']
                self.log_activity(f"   Files Replaced: {exec_data.get('files_replaced', 'Unknown')}")
            
            return payload_data
            
        except Exception as e:
            self.log_activity(f"Failed to decrypt Odyssey payload: {str(e)}", "ERROR")
            return None
    
    def perform_comprehensive_scan(self, directories=None):
        """Perform comprehensive system scan focusing on encrypted files"""
        if directories is None:
            directories = [".", "odyssey_test_environment"]
        
        self.log_activity("Starting comprehensive Odyssey virus scan...", "INFO")
        self.log_activity("Focus: Encrypted files and virus artifacts", "INFO")
        self.log_activity(f"Document format support:", "INFO")
        self.log_activity(f"   Word (.docx): {'YES' if DOCX_AVAILABLE else 'NO'}")
        self.log_activity(f"   PDF (.pdf): {'YES' if PDF_AVAILABLE else 'NO'}")
        self.log_activity(f"   Excel (.xlsx): {'YES' if XLSX_AVAILABLE else 'NO'}")
        self.log_activity(f"   PowerPoint (.pptx): {'YES' if PPTX_AVAILABLE else 'NO'}")
        
        scan_start_time = time.time()
        all_detections = []
        
        for directory in directories:
            if not os.path.exists(directory):
                continue
                
            self.log_activity(f"Scanning directory: {directory}")
            
            # Scan filename patterns (encrypted files priority)
            pattern_detections = self.scan_filename_patterns(directory)
            all_detections.extend(pattern_detections)
            
            # Comprehensive encrypted file scanning
            try:
                for root, dirs, files in os.walk(directory):
                    if self.quarantine_dir in root:
                        continue
                        
                    for filename in files:
                        filepath = os.path.join(root, filename)
                        encrypted_file_detections = self.scan_encrypted_file_comprehensive(filepath)
                        all_detections.extend(encrypted_file_detections)
                        
            except Exception as e:
                self.log_activity(f"Error during comprehensive scan in {directory}: {str(e)}", "ERROR")
            
            # Analyze special files
            marker_data = self.analyze_infection_marker(directory)
            payload_data = self.analyze_encrypted_payload(directory)
        
        # Update statistics
        scan_duration = time.time() - scan_start_time
        self.scan_statistics['scan_duration'] = scan_duration
        self.scan_statistics['threats_detected'] = len(all_detections)
        
        self.log_activity(f"Scan completed in {scan_duration:.2f} seconds", "SUCCESS")
        self.log_activity(f"Scan Statistics:")
        self.log_activity(f"   Files scanned: {self.scan_statistics['files_scanned']}")
        self.log_activity(f"   Encrypted files found: {self.scan_statistics['encrypted_files_found']}")
        self.log_activity(f"   Infected files found: {self.scan_statistics['infected_files_found']}")
        self.log_activity(f"   Artifacts found: {self.scan_statistics['artifacts_found']}")
        self.log_activity(f"   Threats detected: {self.scan_statistics['threats_detected']}")
        
        return all_detections

class OdysseyVirusRemover:
    """
    Virus removal and file recovery system for Odyssey Virus

    Provides comprehensive removal capabilities for encrypted files:
    - Encrypted file decryption and recovery
    - Word document reconstruction from encrypted content
    - Multi-format document restoration from encrypted versions
    - Secure quarantine management
    - Content recovery (original files are deleted)
    - Artifact cleanup
    """
    
    def __init__(self, scanner):
        self.scanner = scanner
        self.crypto_engine = ROT13CryptographyEngine()
        self.document_analyzer = DocumentAnalyzer(self.crypto_engine)
        self.removal_log = []
        self.recovery_statistics = {
            'encrypted_files_processed': 0,
            'infected_files_processed': 0,
            'files_decrypted': 0,
            'files_cleaned': 0,
            'documents_recovered': 0,
            'word_docs_recovered': 0,
            'content_recovered': 0,
            'artifacts_removed': 0,
            'files_quarantined': 0
        }
    
    def log_removal(self, message, level="INFO"):
        """Log removal activity"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = f"[{timestamp}] {message}"
        self.removal_log.append(log_entry)
        self.scanner.log_activity(message, level)
    
    def quarantine_file(self, filepath, reason="Virus detected"):
        """Safely quarantine infected file"""
        try:
            filename = os.path.basename(filepath)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            quarantine_name = f"{timestamp}_{filename}"
            quarantine_path = os.path.join(self.scanner.quarantine_dir, quarantine_name)
            
            # Create quarantine metadata
            metadata = {
                "original_path": filepath,
                "quarantine_time": datetime.now().isoformat(),
                "reason": reason,
                "file_hash": self.scanner.calculate_file_hash(filepath),
                "file_size": os.path.getsize(filepath),
                "encrypted_file": "ODYSSEY_ENCRYPTED" in filename
            }
            
            # Move file to quarantine
            os.rename(filepath, quarantine_path)
            
            # Save metadata
            metadata_path = quarantine_path + ".metadata.json"
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2)
            
            self.recovery_statistics['files_quarantined'] += 1
            self.log_removal(f"Quarantined: {filename} -> {quarantine_name}", "REMOVAL")
            return True
            
        except Exception as e:
            self.log_removal(f"Failed to quarantine {filepath}: {str(e)}", "ERROR")
            return False
    
    def recover_encrypted_word_document(self, encrypted_docx_path, output_path=None):
        """Recover encrypted Word document"""
        if not DOCX_AVAILABLE:
            return False, "python-docx library not available"
        
        try:
            # Analyze the encrypted document
            analysis_result, error = self.document_analyzer.analyze_document(encrypted_docx_path)
            
            if error:
                return False, error
            
            if not analysis_result.get('is_infected', False):
                return False, "Document does not appear to be encrypted"
            
            # Extract text content and decrypt
            encrypted_content = analysis_result['text_content']
            
            # Find the actual encrypted content (skip virus metadata)
            content_lines = encrypted_content.split('\n')
            decrypted_lines = []
            
            for line in content_lines:
                # Skip virus signature lines and metadata
                if not (line.strip().startswith('<!--') or 
                       'ODYSSEY_VIRUS_2025_NTC' in line or
                       line.startswith('Encrypted Document - ROT13') or
                       line.startswith('Educational Note:') or
                       line.startswith('This document has been encrypted') or
                       line.startswith('ROT13 is a simple Caesar cipher') or
                       line.startswith('This is part of a cybersecurity')):
                    # Decrypt the line
                    decrypted_line = self.crypto_engine.rot13_decrypt(line)
                    if decrypted_line.strip():  # Only add non-empty lines
                        decrypted_lines.append(decrypted_line)
            
            decrypted_content = '\n'.join(decrypted_lines)
            
            # Create recovered document
            if output_path is None:
                base_name = os.path.basename(encrypted_docx_path).replace("_ODYSSEY_ENCRYPTED", "")
                name, ext = os.path.splitext(base_name)
                output_path = os.path.join(os.path.dirname(encrypted_docx_path), f"{name}_RECOVERED{ext}")
            
            recovered_doc = Document()
            recovered_doc.add_heading('Recovered Document', 0)
            recovered_doc.add_paragraph('This document has been recovered from ROT13 encryption.')
            recovered_doc.add_paragraph('')
            
            # Add decrypted content
            for paragraph in decrypted_content.split('\n'):
                if paragraph.strip():
                    recovered_doc.add_paragraph(paragraph)
            
            # Add recovery metadata
            recovered_doc.add_paragraph('')
            recovered_doc.add_paragraph('--- Recovery Information ---')
            recovered_doc.add_paragraph(f'Encrypted file: {os.path.basename(encrypted_docx_path)}')
            recovered_doc.add_paragraph(f'Recovery time: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            recovered_doc.add_paragraph('Recovered by: Odyssey Virus Hunter')
            recovered_doc.add_paragraph('NOTE: Original file was deleted by virus')
            
            recovered_doc.save(output_path)
            
            self.recovery_statistics['word_docs_recovered'] += 1
            self.recovery_statistics['content_recovered'] += 1
            self.log_removal(f"Word document recovered: {os.path.basename(encrypted_docx_path)}", "CRYPTO")
            return True, None
            
        except Exception as e:
            return False, f"Error recovering Word document: {str(e)}"
    
    def decrypt_encrypted_text_file(self, encrypted_file_path, output_path=None):
        """Decrypt encrypted text file"""
        try:
            # Determine output path
            if output_path is None:
                base_name = os.path.basename(encrypted_file_path).replace("_ODYSSEY_ENCRYPTED", "")
                if base_name.endswith('.txt'):
                    # Try to restore original extension if possible
                    # Look for original type in the content
                    try:
                        with open(encrypted_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read()
                        
                        # Look for document type indicator
                        for line in content.split('\n'):
                            if line.startswith('Original File:'):
                                original_name = line.split('Original File:')[1].strip()
                                name, ext = os.path.splitext(original_name)
                                output_path = os.path.join(os.path.dirname(encrypted_file_path), f"{name}_RECOVERED{ext}")
                                break
                        
                        if output_path is None:
                            name, ext = os.path.splitext(base_name)
                            output_path = os.path.join(os.path.dirname(encrypted_file_path), f"{name}_RECOVERED{ext}")
                            
                    except:
                        name, ext = os.path.splitext(base_name)
                        output_path = os.path.join(os.path.dirname(encrypted_file_path), f"{name}_RECOVERED{ext}")
                else:
                    name, ext = os.path.splitext(base_name)
                    output_path = os.path.join(os.path.dirname(encrypted_file_path), f"{name}_RECOVERED{ext}")
            
            # Read encrypted file
            with open(encrypted_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                encrypted_content = f.read()
            
            # Extract the actual encrypted content
            lines = encrypted_content.split('\n')
            content_start = 0
            
            # Find where actual content starts (after headers)
            for i, line in enumerate(lines):
                if '=' * 50 in line:
                    content_start = i + 1
                    break
            
            if content_start < len(lines):
                # Extract encrypted content (before virus signatures)
                actual_content = []
                for line in lines[content_start:]:
                    if not line.strip().startswith('<!--'):
                        actual_content.append(line)
                
                # Remove trailing virus signatures
                while actual_content and actual_content[-1].strip().startswith('<!--'):
                    actual_content.pop()
                
                encrypted_text = '\n'.join(actual_content)
                decrypted_content = self.crypto_engine.rot13_decrypt(encrypted_text)
                
                # Write decrypted file
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(decrypted_content)
                    f.write(f"\n\n--- Recovery Information ---")
                    f.write(f"\nEncrypted file: {os.path.basename(encrypted_file_path)}")
                    f.write(f"\nRecovery time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    f.write(f"\nRecovered by: Odyssey Virus Hunter")
                    f.write(f"\nNOTE: Original file was deleted by virus")
                
                self.recovery_statistics['files_decrypted'] += 1
                self.recovery_statistics['content_recovered'] += 1
                self.log_removal(f"Decrypted: {os.path.basename(encrypted_file_path)} -> {os.path.basename(output_path)}", "CRYPTO")
                return True
            else:
                self.log_removal(f"No encrypted content found in {encrypted_file_path}", "WARNING")
                return False
            
        except Exception as e:
            self.log_removal(f"Failed to decrypt {encrypted_file_path}: {str(e)}", "ERROR")
            return False
    
    def process_encryption_manifest(self, directory):
        """Process Odyssey virus encryption manifest for recovery"""
        manifest_path = os.path.join(directory, "odyssey_encryption_manifest.json")
        
        if not os.path.exists(manifest_path):
            self.log_removal("No encryption manifest found")
            return 0
        
        try:
            with open(manifest_path, 'r', encoding='utf-8') as f:
                manifest_data = json.load(f)
            
            self.log_removal(f"Processing encryption manifest: {len(manifest_data)} entries")
            
            decrypted_count = 0
            
            for encrypted_filename, file_info in manifest_data.items():
                encrypted_path = os.path.join(directory, encrypted_filename)
                
                if os.path.exists(encrypted_path):
                    self.recovery_statistics['encrypted_files_processed'] += 1
                    original_filename = file_info.get('original_filename', 'unknown')
                    original_type = file_info.get('original_type', '.txt')
                    
                    self.log_removal(f"Processing encrypted file: {encrypted_filename}")
                    self.log_removal(f"  Original: {original_filename} ({original_type})")
                    
                    # Handle Word documents specially
                    if encrypted_filename.endswith('_ODYSSEY_ENCRYPTED.docx') and DOCX_AVAILABLE:
                        success, error = self.recover_encrypted_word_document(encrypted_path)
                        if success:
                            decrypted_count += 1
                            # Remove encrypted version
                            try:
                                os.remove(encrypted_path)
                                self.log_removal(f"Removed encrypted Word document: {encrypted_filename}")
                            except Exception as e:
                                self.log_removal(f"Could not remove {encrypted_filename}: {str(e)}", "WARNING")
                        else:
                            self.log_removal(f"Failed to recover Word document: {error}", "ERROR")
                    
                    else:
                        # Handle other document types (encrypted as text)
                        success = self.decrypt_encrypted_text_file(encrypted_path)
                        if success:
                            decrypted_count += 1
                            self.recovery_statistics['documents_recovered'] += 1
                            self.log_removal(f"Document recovered: {original_filename} ({original_type})", "CRYPTO")
                            
                            # Remove encrypted version
                            try:
                                os.remove(encrypted_path)
                                self.log_removal(f"Removed encrypted file: {encrypted_filename}")
                            except Exception as e:
                                self.log_removal(f"Could not remove {encrypted_filename}: {str(e)}", "WARNING")
                        else:
                            self.log_removal(f"Failed to recover {encrypted_filename}", "ERROR")
                else:
                    self.log_removal(f"Encrypted file not found: {encrypted_filename}", "WARNING")
            
            # Remove the manifest
            try:
                os.remove(manifest_path)
                self.log_removal("Removed encryption manifest")
            except Exception as e:
                self.log_removal(f"Could not remove manifest: {str(e)}", "WARNING")
            
            return decrypted_count
            
        except Exception as e:
            self.log_removal(f"Error processing encryption manifest: {str(e)}", "ERROR")
            return 0
    
    def restore_locked_files(self, directory):
        """Restore files that were temporarily locked by virus (encrypted files only)"""
        restored_count = 0
        
        try:
            for root, dirs, files in os.walk(directory):
                if self.scanner.quarantine_dir in root:
                    continue
                    
                for filename in files:
                    if filename.startswith("ODYSSEY_LOCKED_"):
                        original_name = filename[15:]  # Remove "ODYSSEY_LOCKED_" prefix
                        original_path = os.path.join(root, original_name)
                        locked_path = os.path.join(root, filename)
                        
                        # Only restore if original doesn't exist
                        if not os.path.exists(original_path):
                            try:
                                os.rename(locked_path, original_path)
                                restored_count += 1
                                self.log_removal(f"Restored locked file: {filename} -> {original_name}", "SUCCESS")
                            except Exception as e:
                                self.log_removal(f"Failed to restore {filename}: {str(e)}", "ERROR")
                        else:
                            # Original exists, just remove the locked version
                            try:
                                os.remove(locked_path)
                                self.log_removal(f"Removed duplicate locked file: {filename}")
                            except Exception as e:
                                self.log_removal(f"Could not remove locked file {filename}: {str(e)}", "WARNING")
                                
        except Exception as e:
            self.log_removal(f"Error during locked file restoration: {str(e)}", "ERROR")
        
        return restored_count
    
    def clean_virus_artifacts(self, directory):
        """Remove all Odyssey virus artifacts"""
        artifacts_removed = 0
        
        # Known Odyssey virus artifacts
        artifact_patterns = [
            "odyssey_activity.log",
            ".odyssey_infection_marker",
            "odyssey_encrypted_payload.dat",
            "odyssey_encryption_manifest.json",
            "*ODYSSEY_LOCKED_*"
        ]
        
        try:
            for root, dirs, files in os.walk(directory):
                if self.scanner.quarantine_dir in root:
                    continue
                    
                for filename in files:
                    filepath = os.path.join(root, filename)
                    
                    # Check if file matches virus artifact patterns
                    for pattern in artifact_patterns:
                        if fnmatch.fnmatch(filename, pattern):
                            try:
                                os.remove(filepath)
                                self.recovery_statistics['artifacts_removed'] += 1
                                self.log_removal(f"Removed artifact: {filename}", "REMOVAL")
                                artifacts_removed += 1
                            except Exception as e:
                                self.log_removal(f"Failed to remove {filename}: {str(e)}", "ERROR")
                            break
                            
        except Exception as e:
            self.log_removal(f"Error during artifact cleanup: {str(e)}", "ERROR")
        
        return artifacts_removed

    def process_infection_manifest(self, directory):
        """Process infection manifest and clean infected files"""
        manifest_path = os.path.join(directory, "odyssey_infection_manifest.json")
        
        if not os.path.exists(manifest_path):
            self.log_removal("No infection manifest found")
            return 0
        
        try:
            with open(manifest_path, 'r', encoding='utf-8') as f:
                manifest_data = json.load(f)
            
            self.log_removal(f"Processing infection manifest: {len(manifest_data)} entries")
            
            cleaned_count = 0
            
            for filename, file_info in manifest_data.items():
                file_path = os.path.join(directory, filename)
                
                if os.path.exists(file_path):
                    self.recovery_statistics['infected_files_processed'] += 1
                    self.log_removal(f"Cleaning infected file: {filename}")
                    
                    success = self.clean_infected_file(file_path)
                    if success:
                        cleaned_count += 1
                        self.log_removal(f"Successfully cleaned: {filename}", "SUCCESS")
                    else:
                        self.log_removal(f"Failed to clean: {filename}", "ERROR")
                else:
                    self.log_removal(f"Infected file not found: {filename}", "WARNING")
            
            # Remove the infection manifest
            try:
                os.remove(manifest_path)
                self.log_removal("Removed infection manifest")
            except Exception as e:
                self.log_removal(f"Could not remove infection manifest: {str(e)}", "WARNING")
            
            return cleaned_count
            
        except Exception as e:
            self.log_removal(f"Error processing infection manifest: {str(e)}", "ERROR")
            return 0

    def clean_infected_file(self, infected_file_path):
        """Clean virus infections from file"""
        try:
            filename = os.path.basename(infected_file_path)
            
            # Check for backup file
            backup_path = infected_file_path + ".backup"
            
            if os.path.exists(backup_path):
                # Restore from backup
                try:
                    with open(backup_path, 'r', encoding='utf-8') as f:
                        original_content = f.read()
                    
                    with open(infected_file_path, 'w', encoding='utf-8') as f:
                        f.write(original_content)
                    
                    # Remove backup
                    os.remove(backup_path)
                    
                    self.recovery_statistics['files_cleaned'] += 1
                    self.log_removal(f"Restored from backup: {filename}", "SUCCESS")
                    return True
                    
                except Exception as e:
                    self.log_removal(f"Failed to restore from backup {filename}: {str(e)}", "ERROR")
                    return False
            else:
                # Manual cleaning - remove virus signatures
                try:
                    with open(infected_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        infected_content = f.read()
                    
                    # Remove known virus signatures
                    cleaned_content = infected_content
                    
                    # Remove infection markers
                    infection_patterns = [
                        "ODYSSEY_VIRUS_2025_NTC INFECTION MARKER",
                        "Educational Malware Injection",
                        "Virus Code Injection",
                        "Script Infection",
                        "INFECTED BY ODYSSEY_VIRUS_2025_NTC",
                        "Infection Time:",
                        "Educational Purpose: Malware Behavior Simulation",
                        "Institution: National Teachers College",
                        "Educational virus injection completed",
                        "Educational Virus Injection",
                        "Infection performed for educational cybersecurity research",
                        "National Teachers College - Security Analysis Project",
                        "Educational Purpose Only",
                        "Educational Malware Simulation"
                    ]
                    
                    # Remove lines containing infection patterns
                    lines = cleaned_content.split('\n')
                    clean_lines = []
                    
                    for line in lines:
                        should_remove = False
                        for pattern in infection_patterns:
                            if pattern in line:
                                should_remove = True
                                break
                        
                        # Also remove HTML/CSS comments and Python comments with virus signatures
                        if (line.strip().startswith('<!--') and 'ODYSSEY' in line) or \
                        (line.strip().startswith('#') and 'ODYSSEY' in line) or \
                        (line.strip().startswith('/*') and 'ODYSSEY' in line) or \
                        (line.strip().startswith('//') and 'ODYSSEY' in line):
                            should_remove = True
                        
                        if not should_remove:
                            clean_lines.append(line)
                    
                    cleaned_content = '\n'.join(clean_lines)
                    
                    # Remove trailing infection markers
                    while cleaned_content.endswith('\n# INFECTED BY ODYSSEY_VIRUS_2025_NTC #') or \
                        cleaned_content.endswith('# INFECTED BY ODYSSEY_VIRUS_2025_NTC #'):
                        cleaned_content = cleaned_content.rsplit('# INFECTED BY ODYSSEY_VIRUS_2025_NTC #', 1)[0].rstrip()
                    
                    # Write cleaned content
                    with open(infected_file_path, 'w', encoding='utf-8') as f:
                        f.write(cleaned_content)
                        f.write(f"\n\n# File cleaned by Odyssey Antivirus")
                        f.write(f"\n# Cleaning time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                        f.write(f"\n# Virus signatures removed")
                    
                    self.recovery_statistics['files_cleaned'] += 1
                    self.log_removal(f"Manually cleaned: {filename}", "SUCCESS")
                    return True
                    
                except Exception as e:
                    self.log_removal(f"Failed to clean {filename}: {str(e)}", "ERROR")
                    return False
                    
        except Exception as e:
            self.log_removal(f"Error cleaning infected file {infected_file_path}: {str(e)}", "ERROR")
            return False
    
    def perform_complete_removal(self, target_directory="odyssey_test_environment"):
        """Perform complete Odyssey virus removal and recovery"""
        self.log_removal("Starting comprehensive Odyssey virus removal...", "REMOVAL")
        self.log_removal("Focus: Encrypted file recovery and artifact cleanup", "REMOVAL")
        
        removal_start_time = time.time()
        
        if not os.path.exists(target_directory):
            self.log_removal(f"Target directory not found: {target_directory}", "WARNING")
            return False
        
        # Step 1: Process infection manifest and clean infected files
        self.log_removal("Step 1: Processing infection manifest and cleaning infected files...")
        cleaned_count = self.process_infection_manifest(target_directory)# Fixed call
        
        # Step 2: Process encryption manifest and decrypt files
        self.log_removal("Step 2: Processing encryption manifest and recovering encrypted files...")
        decrypted_count = self.process_encryption_manifest(target_directory)
        
        # Step 3: Restore locked files (encrypted files only)
        self.log_removal("Step 3: Restoring temporarily locked files...")
        restored_count = self.restore_locked_files(target_directory)
        
        # Step 4: Clean up virus artifacts
        self.log_removal("Step 4: Cleaning virus artifacts...")
        artifacts_removed = self.clean_virus_artifacts(target_directory)
        
        # Calculate removal duration
        removal_duration = time.time() - removal_start_time
        
        # Log summary
        self.log_removal(f"Removal completed in {removal_duration:.2f} seconds", "SUCCESS")
        self.log_removal(f"Recovery Statistics:")
        self.log_removal(f"   Infected files processed: {self.recovery_statistics['infected_files_processed']}")
        self.log_removal(f"   Files cleaned: {self.recovery_statistics['files_cleaned']}")
        self.log_removal(f"   Encrypted files processed: {self.recovery_statistics['encrypted_files_processed']}")
        self.log_removal(f"   Files decrypted: {self.recovery_statistics['files_decrypted']}")
        self.log_removal(f"   Documents recovered: {self.recovery_statistics['documents_recovered']}")
        self.log_removal(f"   Word documents recovered: {self.recovery_statistics['word_docs_recovered']}")
        self.log_removal(f"   Content recovered: {self.recovery_statistics['content_recovered']}")
        self.log_removal(f"   Files restored: {restored_count}")
        self.log_removal(f"   Artifacts removed: {artifacts_removed}")
        self.log_removal("NOTE: Original files were deleted by virus - only recovered content and cleaned files available")
        
        return True


class AntivirusGUI:
    """
    GUI for the Odyssey Virus Hunter
    
    Features:
    - Real-time scanning progress focused on encrypted files
    - Multi-document threat analysis
    - Document format support indicators
    - Interactive removal options for encrypted content
    - Comprehensive reporting with recovery focus
    """
    
    def __init__(self):
        print("DEBUG: Initializing AntivirusGUI...")
        
        try:
            self.root = tk.Tk()
            self.root.title("Odyssey Virus Hunter")
            self.root.geometry("1200x800")
            self.root.configure(bg='#f0f0f0')
            print("DEBUG: tkinter root window created successfully")
            
            self.scanner = AdvancedVirusScanner()
            self.remover = OdysseyVirusRemover(self.scanner)
            self.detections = []
            self.scan_in_progress = False
            
            print("DEBUG: Starting GUI setup...")
            self.setup_gui()
            print("DEBUG: GUI setup completed successfully")
            
        except Exception as e:
            print(f"ERROR: Failed to initialize GUI: {e}")
            raise

    def setup_gui(self):
        """Setup GUI components"""
        try:
            # Create main notebook for tabs
            self.notebook = ttk.Notebook(self.root)
            self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            # Virus Scanner tab
            self.scanner_frame = tk.Frame(self.notebook, bg='#f0f0f0')
            self.notebook.add(self.scanner_frame, text="Encrypted File Scanner")
            self.setup_scanner_tab()
            
            # Detection Results tab
            self.results_frame = tk.Frame(self.notebook, bg='#f0f0f0')
            self.notebook.add(self.results_frame, text="Detection Results")
            self.setup_results_tab()
            
            # Recovery Analysis tab
            self.analysis_frame = tk.Frame(self.notebook, bg='#f0f0f0')
            self.notebook.add(self.analysis_frame, text="Recovery Analysis")
            self.setup_analysis_tab()

            # Quarantine tab
            self.quarantine_frame = tk.Frame(self.notebook, bg='#f0f0f0')
            self.notebook.add(self.quarantine_frame, text="Quarantine Management")
            self.setup_quarantine_tab()
            
        except Exception as e:
            print(f"ERROR: Failed to setup GUI: {e}")
            raise
        
    def setup_scanner_tab(self):
        """Setup the scanner interface with improved layout"""
        
        # Main container with better space management
        main_container = tk.Frame(self.scanner_frame, bg='#f0f0f0')
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Top section - titles and notices
        top_section = tk.Frame(main_container, bg='#f0f0f0')
        top_section.pack(fill=tk.X, pady=(0, 10))
        
        # Title
        title_label = tk.Label(top_section,
                             text="Odyssey Virus Hunter",
                             font=("Arial", 18, "bold"),
                             bg='#f0f0f0', fg='#333')
        title_label.pack(pady=(10, 5))
        
        subtitle_label = tk.Label(top_section,
                                text="Encrypted File Detection & Recovery System",
                                font=("Arial", 12),
                                bg='#f0f0f0', fg='#666')
        subtitle_label.pack(pady=(0, 10))
        
        # Compact Important Notice
        behavior_frame = tk.LabelFrame(top_section, text="Important Notice", 
                                     bg='#fff3cd', font=("Arial", 9, "bold"))
        behavior_frame.pack(fill=tk.X, pady=(0, 5))
        
        # More concise notice text
        behavior_text = """⚠️  Original files DELETED by virus • Only encrypted files remain • Focus on recovery"""
        
        behavior_label = tk.Label(behavior_frame, text=behavior_text.strip(),
                                bg='#fff3cd', font=("Arial", 9), justify=tk.CENTER, fg='#856404')
        behavior_label.pack(padx=10, pady=8)
        
        # Compact Document support status
        support_frame = tk.LabelFrame(top_section, text="Document Support", 
                                    bg='#f0f0f0', font=("Arial", 9, "bold"))
        support_frame.pack(fill=tk.X, pady=(0, 5))
        
        # Create a more compact layout using grid
        support_container = tk.Frame(support_frame, bg='#f0f0f0')
        support_container.pack(padx=10, pady=5)
        
        # First row
        tk.Label(support_container, text=f"Word: {'✓' if DOCX_AVAILABLE else '✗'}", 
                 bg='#f0f0f0', font=("Arial", 8), fg='green' if DOCX_AVAILABLE else 'red').grid(row=0, column=0, padx=5)
        tk.Label(support_container, text=f"PDF: {'✓' if PDF_AVAILABLE else '✗'}", 
                 bg='#f0f0f0', font=("Arial", 8), fg='green' if PDF_AVAILABLE else 'red').grid(row=0, column=1, padx=5)
        tk.Label(support_container, text=f"Excel: {'✓' if XLSX_AVAILABLE else '✗'}", 
                 bg='#f0f0f0', font=("Arial", 8), fg='green' if XLSX_AVAILABLE else 'red').grid(row=0, column=2, padx=5)
        tk.Label(support_container, text=f"PowerPoint: {'✓' if PPTX_AVAILABLE else '✗'}", 
                 bg='#f0f0f0', font=("Arial", 8), fg='green' if PPTX_AVAILABLE else 'red').grid(row=0, column=3, padx=5)
        tk.Label(support_container, text="Text: ✓", 
                 bg='#f0f0f0', font=("Arial", 8), fg='green').grid(row=0, column=4, padx=5)
        
        # Control buttons section
        control_frame = tk.Frame(top_section, bg='#f0f0f0')
        control_frame.pack(pady=10)

        # Scan button
        self.scan_button = tk.Button(control_frame,
                                   text="Start Encrypted File Scan",
                                   command=self.start_scan_thread,
                                   bg="#4CAF50", fg="white",
                                   font=("Arial", 11, "bold"),
                                   padx=25, pady=12,
                                   cursor="hand2")
        self.scan_button.pack(side=tk.LEFT, padx=8)
        
        # Quick scan button
        quick_scan_button = tk.Button(control_frame,
                                    text="Quick Scan",
                                    command=self.quick_scan,
                                    bg="#2196F3", fg="white",
                                    font=("Arial", 10),
                                    padx=15, pady=8,
                                    cursor="hand2")
        quick_scan_button.pack(side=tk.LEFT, padx=8)
        
        # Middle section - Progress and Statistics (using PanedWindow for resizable layout)
        middle_paned = tk.PanedWindow(main_container, orient=tk.VERTICAL, bg='#f0f0f0', sashrelief=tk.RAISED, sashwidth=3)
        middle_paned.pack(fill=tk.BOTH, expand=True)
        
        # Progress frame
        progress_frame = tk.LabelFrame(middle_paned, text="Scan Progress", bg='#f0f0f0', font=("Arial", 9, "bold"))
        middle_paned.add(progress_frame, height=80, minsize=60)

        self.progress_var = tk.StringVar(value="Ready for encrypted file scan")
        self.progress_label = tk.Label(progress_frame, textvariable=self.progress_var,
                                     bg='#f0f0f0', fg='#666', font=("Arial", 9))
        self.progress_label.pack(anchor=tk.W, padx=10, pady=5)
        
        self.progress_bar = ttk.Progressbar(progress_frame, mode='indeterminate')
        self.progress_bar.pack(fill=tk.X, padx=10, pady=5)
        
        # Statistics display with better space allocation
        stats_frame = tk.LabelFrame(middle_paned, text="Scan Statistics", bg='#f0f0f0', font=("Arial", 9, "bold"))
        middle_paned.add(stats_frame, minsize=120)
        
        # Create a frame with scrollable text for statistics
        stats_container = tk.Frame(stats_frame, bg='#f0f0f0')
        stats_container.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.stats_text = scrolledtext.ScrolledText(stats_container, 
                                                   height=8, 
                                                   font=("Courier", 9), 
                                                   bg='#ffffff',
                                                   wrap=tk.WORD)
        self.stats_text.pack(fill=tk.BOTH, expand=True)
        
        # Console output with flexible sizing
        console_frame = tk.LabelFrame(middle_paned, text="Scanner Console", bg='#f0f0f0', font=("Arial", 9, "bold"))
        middle_paned.add(console_frame, minsize=150)
        
        self.console_text = scrolledtext.ScrolledText(console_frame,
                                                    font=("Courier", 9),
                                                    bg='#1e1e1e', fg='#00ff00',
                                                    insertbackground='#00ff00',
                                                    wrap=tk.WORD)
        self.console_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Set initial pane sizes (progress:stats:console = 1:2:3 ratio approximately)
        self.scanner_frame.after(100, lambda: middle_paned.paneconfigure(progress_frame, height=80))
        self.scanner_frame.after(100, lambda: middle_paned.paneconfigure(stats_frame, height=160))
        self.scanner_frame.after(100, lambda: middle_paned.paneconfigure(console_frame, height=240))
        
    def setup_results_tab(self):
        """Setup the results analysis tab"""
        # Results header
        header_frame = tk.Frame(self.results_frame, bg='#f0f0f0')
        header_frame.pack(fill=tk.X, padx=20, pady=10)

        tk.Label(header_frame, text="Encrypted File Detection Results",
                font=("Arial", 14, "bold"), bg='#f0f0f0').pack(anchor=tk.W)

        # Action buttons
        action_frame = tk.Frame(self.results_frame, bg='#f0f0f0')
        action_frame.pack(fill=tk.X, padx=20, pady=10)
        
        self.recover_button = tk.Button(action_frame,
                                      text="Recover & Clean All Files",
                                      command=self.recover_encrypted_files,
                                      bg="#4CAF50", fg="white",
                                      font=("Arial", 11, "bold"),
                                      padx=20, pady=10,
                                      cursor="hand2")
        self.recover_button.pack(side=tk.LEFT, padx=5)
        
        self.quarantine_button = tk.Button(action_frame,
                                         text="Quarantine Threats",
                                         command=self.quarantine_threats,
                                         bg="#FF9800", fg="white",
                                         font=("Arial", 11),
                                         padx=20, pady=10,
                                         cursor="hand2")
        self.quarantine_button.pack(side=tk.LEFT, padx=5)

        # Results display
        self.results_text = scrolledtext.ScrolledText(self.results_frame,
                                                    height=25,
                                                    font=("Courier", 10),
                                                    bg='#ffffff')
        self.results_text.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
    def setup_analysis_tab(self):
        """Setup the recovery analysis tab"""
        # Analysis header
        header_frame = tk.Frame(self.analysis_frame, bg='#f0f0f0')
        header_frame.pack(fill=tk.X, padx=20, pady=10)

        tk.Label(header_frame, text="Encrypted File Recovery Analysis",
                font=("Arial", 14, "bold"), bg='#f0f0f0').pack(anchor=tk.W)
        
        # Analysis display
        self.analysis_text = scrolledtext.ScrolledText(self.analysis_frame,
                                                     height=25,
                                                     font=("Courier", 10),
                                                     bg='#ffffff')
        self.analysis_text.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
    def setup_quarantine_tab(self):
        """Setup the quarantine management tab"""
        # Quarantine header
        header_frame = tk.Frame(self.quarantine_frame, bg='#f0f0f0')
        header_frame.pack(fill=tk.X, padx=20, pady=10)

        tk.Label(header_frame, text="Quarantine Management",
                font=("Arial", 14, "bold"), bg='#f0f0f0').pack(anchor=tk.W)

        # Quarantine info
        self.quarantine_text = scrolledtext.ScrolledText(self.quarantine_frame,
                                                       height=20,
                                                       font=("Courier", 10),
                                                       bg='#ffffff')
        self.quarantine_text.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # Update quarantine info
        self.update_quarantine_info()
        
    def log_to_console(self, message):
        """Log message to GUI console"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        formatted_message = f"[{timestamp}] {message}\n"
        
        self.console_text.insert(tk.END, formatted_message)
        self.console_text.see(tk.END)
        self.root.update()
        
    def update_progress(self, message):
        """Update progress display"""
        self.progress_var.set(message)
        self.root.update()
        
    def start_scan_thread(self):
        """Start scan in separate thread to prevent GUI freezing"""
        if self.scan_in_progress:
            return
            
        self.scan_in_progress = True
        self.scan_button.config(state=tk.DISABLED, text="Scanning...")
        self.progress_bar.start()
        
        # Start scan in separate thread
        scan_thread = threading.Thread(target=self.perform_scan)
        scan_thread.daemon = True
        scan_thread.start()
        
    def perform_scan(self):
        """Perform the actual scan"""
        try:
            self.console_text.delete(1.0, tk.END)
            self.log_to_console("Odyssey Virus Hunter Starting...")
            
            # Initialize quarantine
            self.update_progress("Initializing quarantine system...")
            self.scanner.initialize_quarantine_system()
            
            # Override scanner's log_activity to also log to GUI
            original_log = self.scanner.log_activity
            
            def gui_log(message, level="INFO"):
                original_log(message, level)
                self.log_to_console(message)
                
            self.scanner.log_activity = gui_log

            # Perform scan
            self.update_progress("Scanning for encrypted files and artifacts...")
            self.detections = self.scanner.perform_comprehensive_scan()
            
            # Update results tab
            self.update_results_display()
            
            # Update statistics
            self.update_statistics_display()
            
            # Update analysis
            self.update_analysis_display()

            self.log_to_console(f"\nScan completed: {len(self.detections)} threats detected")

        except Exception as e:
            self.log_to_console(f"Scan error: {str(e)}")

        finally:
            # Reset UI
            self.scan_in_progress = False
            self.scan_button.config(state=tk.NORMAL, text="Start Encrypted File Scan")
            self.progress_bar.stop()
            self.update_progress("Scan completed")

    def quick_scan(self):
        """Perform quick scan of current directory only"""
        self.log_to_console("Starting quick encrypted file scan...")
        self.detections = self.scanner.perform_comprehensive_scan(["."])
        self.update_results_display()
        self.update_analysis_display()
        self.log_to_console(f"Quick scan completed: {len(self.detections)} threats detected")

    def update_results_display(self):
        """Update the results tab with detection information"""
        self.results_text.delete(1.0, tk.END)
        
        if not self.detections:
            self.results_text.insert(tk.END, "No encrypted files or threats detected. System appears clean.\n")
            return

        self.results_text.insert(tk.END, f"ENCRYPTED FILE DETECTION REPORT\n")
        self.results_text.insert(tk.END, f"=" * 60 + "\n")
        self.results_text.insert(tk.END, f"Total threats detected: {len(self.detections)}\n")
        self.results_text.insert(tk.END, f"Scan time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")

        # Group by detection type
        by_type = {}
        for detection in self.detections:
            det_type = detection['type']
            if det_type not in by_type:
                by_type[det_type] = []
            by_type[det_type].append(detection)
        
        for det_type, detections in by_type.items():
            self.results_text.insert(tk.END, f"--- {det_type.upper().replace('_', ' ')} ({len(detections)}) ---\n")
            
            for i, detection in enumerate(detections, 1):
                self.results_text.insert(tk.END, f"\n{i}. File: {detection['file']}\n")
                
                if 'document_type' in detection:
                    self.results_text.insert(tk.END, f"   Document Type: {detection['document_type']}\n")
                
                if 'virus_indicators' in detection:
                    self.results_text.insert(tk.END, f"   Virus Indicators: {len(detection['virus_indicators'])}\n")
                    for indicator in detection['virus_indicators'][:3]:  # Show first 3
                        self.results_text.insert(tk.END, f"     • {indicator}\n")
                
                if 'virus_info' in detection:
                    info = detection['virus_info']
                    self.results_text.insert(tk.END, f"   Virus: {info['name']}\n")
                    self.results_text.insert(tk.END, f"   Risk Level: {info['risk_level']}\n")
                    
                if 'description' in detection:
                    self.results_text.insert(tk.END, f"   Description: {detection['description']}\n")
                
                if detection.get('encrypted', False):
                    self.results_text.insert(tk.END, f"   Status: ENCRYPTED FILE (recoverable)\n")
                
                self.results_text.insert(tk.END, f"   Confidence: {detection.get('confidence', 'MEDIUM')}\n")
                self.results_text.insert(tk.END, f"   Detection Time: {detection.get('detection_time', 'Unknown')}\n")
            
            self.results_text.insert(tk.END, "\n")
            
    def update_statistics_display(self):
        """Update scan statistics display"""
        stats = self.scanner.scan_statistics
        
        stats_text = f"""
Files Scanned:          {stats['files_scanned']}
Encrypted Files Found:  {stats['encrypted_files_found']}
Infected Files Found:   {stats['infected_files_found']}
Artifacts Found:        {stats['artifacts_found']}
Threats Detected:       {stats['threats_detected']}
Scan Duration:          {stats['scan_duration']:.2f} seconds
Files Quarantined:      {stats.get('files_quarantined', 0)}
Files Cleaned:          {stats.get('files_cleaned', 0)}
Last Scan:             {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

Document Types Scanned:
"""
        
        if stats['document_types_scanned']:
            for doc_type, count in stats['document_types_scanned'].items():
                stats_text += f"  {doc_type}: {count}\n"
        else:
            stats_text += "  No documents processed\n"
        
        self.stats_text.delete(1.0, tk.END)
        self.stats_text.insert(tk.END, stats_text.strip())
        
    def update_analysis_display(self):
        """Update recovery analysis display"""
        self.analysis_text.delete(1.0, tk.END)
        
        stats = self.scanner.scan_statistics
        
        analysis_report = f"""
ENCRYPTED FILE RECOVERY ANALYSIS REPORT
{'=' * 60}

SCAN OVERVIEW:
Files Scanned: {stats['files_scanned']}
Encrypted Files Found: {stats['encrypted_files_found']}
Infected Files Found: {stats['infected_files_found']}
Artifacts Found: {stats['artifacts_found']}
Threats Detected: {stats['threats_detected']}
Scan Duration: {stats['scan_duration']:.2f} seconds

DOCUMENT FORMAT SUPPORT:
Word Document Support: {'YES' if DOCX_AVAILABLE else 'NO'}
PDF Support: {'YES' if PDF_AVAILABLE else 'NO'}
Excel Support: {'YES' if XLSX_AVAILABLE else 'NO'}
PowerPoint Support: {'YES' if PPTX_AVAILABLE else 'NO'}

ENCRYPTED FILES ANALYSIS:
"""
        
        # Count encrypted and infected files by type
        encrypted_by_type = {}
        infected_by_type = {}
        for detection in self.detections:
            if detection.get('type') in ['encrypted_file', 'encrypted_document_infection']:
                filename = os.path.basename(detection['file'])
                if filename.endswith('.docx'):
                    file_type = 'Word Document'
                elif filename.endswith('.txt'):
                    file_type = 'Text File'
                else:
                    file_type = 'Other'
                    
                if file_type in encrypted_by_type:
                    encrypted_by_type[file_type] += 1
                else:
                    encrypted_by_type[file_type] = 1
            
            elif detection.get('type') == 'infected_file':
                filename = os.path.basename(detection['file'])
                file_extension = os.path.splitext(filename)[1].lower()
                if file_extension == '.py':
                    file_type = 'Python Script'
                elif file_extension == '.js':
                    file_type = 'JavaScript File'
                elif file_extension in ['.txt', '.md']:
                    file_type = 'Text File'
                elif file_extension == '.csv':
                    file_type = 'CSV File'
                else:
                    file_type = 'Other'
                    
                if file_type in infected_by_type:
                    infected_by_type[file_type] += 1
                else:
                    infected_by_type[file_type] = 1
        
        if encrypted_by_type:
            analysis_report += "ENCRYPTED FILES:\n"
            for file_type, count in encrypted_by_type.items():
                analysis_report += f"• {file_type}: {count} encrypted files\n"
        
        if infected_by_type:
            analysis_report += "INFECTED FILES:\n"
            for file_type, count in infected_by_type.items():
                analysis_report += f"• {file_type}: {count} infected files\n"
        
        if not encrypted_by_type and not infected_by_type:
            analysis_report += "• No encrypted or infected files found\n"
        
        analysis_report += f"""

RECOVERY POTENTIAL:
Detection Rate: {(stats['threats_detected'] / max(stats['files_scanned'], 1) * 100):.1f}%
Encrypted File Rate: {(stats['encrypted_files_found'] / max(stats['files_scanned'], 1) * 100):.1f}%
Infected File Rate: {(stats['infected_files_found'] / max(stats['files_scanned'], 1) * 100):.1f}%

THREAT BREAKDOWN:
"""
        
        # Add threat type breakdown
        threat_types = {}
        for detection in self.detections:
            det_type = detection['type']
            if det_type in threat_types:
                threat_types[det_type] += 1
            else:
                threat_types[det_type] = 1
        
        for threat_type, count in threat_types.items():
            analysis_report += f"• {threat_type.replace('_', ' ').title()}: {count}\n"
        
        analysis_report += f"""

RECOVERY NOTES:
• Original files were deleted by the virus
• Encrypted files can be decrypted using ROT13
• Infected files can be cleaned by removing virus signatures
• Backup files may be available for infected file restoration
• Recovered files will be marked with recovery metadata
• Cleaned files will be marked with cleaning metadata
"""
        
        self.analysis_text.insert(tk.END, analysis_report.strip())
        
    def recover_encrypted_files(self):
        """Recover all detected encrypted files and clean infected files"""
        if not self.detections:
            messagebox.showinfo("No Threats", "No encrypted files or infected files detected to process.")
            return
            
        # Count encrypted and infected files
        encrypted_files = [d for d in self.detections if d.get('type') in ['encrypted_file', 'encrypted_document_infection']]
        infected_files = [d for d in self.detections if d.get('type') == 'infected_file']
        
        if not encrypted_files and not infected_files:
            messagebox.showinfo("No Files to Process", "No encrypted or infected files found to process.")
            return

        result = messagebox.askyesno("Confirm Recovery & Cleaning",
                                   f"Process {len(encrypted_files)} encrypted files and {len(infected_files)} infected files?\n\n"
                                   f"This will:\n"
                                   f"• Decrypt encrypted files and create recovered versions\n"
                                   f"• Clean infected files by removing virus signatures\n"
                                   f"• Restore from backups where available\n\n"
                                   f"Note: Original files were deleted by the virus.")
        if not result:
            return
            
        # Switch to scanner tab to show progress
        self.notebook.select(0)

        self.log_to_console("\nStarting file recovery and cleaning process...")
        
        # Override remover's log to also log to GUI
        original_log = self.remover.log_removal
        
        def gui_removal_log(message, level="INFO"):
            original_log(message, level)
            self.log_to_console(message)
            
        self.remover.log_removal = gui_removal_log

        # Perform recovery and cleaning
        success = self.remover.perform_complete_removal()
        
        if success:
            self.log_to_console("\nFile recovery and cleaning completed successfully!")

            # Update statistics
            recovery_stats = self.remover.recovery_statistics
            self.log_to_console(f"\nRecovery & Cleaning Statistics:")
            self.log_to_console(f"   Infected Files Processed: {recovery_stats['infected_files_processed']}")
            self.log_to_console(f"   Files Cleaned: {recovery_stats['files_cleaned']}")
            self.log_to_console(f"   Encrypted Files Processed: {recovery_stats['encrypted_files_processed']}")
            self.log_to_console(f"   Files Decrypted: {recovery_stats['files_decrypted']}")
            self.log_to_console(f"   Documents Recovered: {recovery_stats['documents_recovered']}")
            self.log_to_console(f"   Word Documents Recovered: {recovery_stats['word_docs_recovered']}")
            self.log_to_console(f"   Content Recovered: {recovery_stats['content_recovered']}")
            self.log_to_console(f"   Artifacts Removed: {recovery_stats['artifacts_removed']}")
            
            # Clear detections
            self.detections = []
            self.update_results_display()
            self.update_analysis_display()

            messagebox.showinfo("Recovery & Cleaning Complete",
                              "All files have been successfully processed!\n\n"
                              f"Infected files processed: {recovery_stats['infected_files_processed']}\n"
                              f"Files cleaned: {recovery_stats['files_cleaned']}\n"
                              f"Encrypted files processed: {recovery_stats['encrypted_files_processed']}\n"
                              f"Files decrypted: {recovery_stats['files_decrypted']}\n"
                              f"Documents recovered: {recovery_stats['documents_recovered']}\n"
                              f"Word docs recovered: {recovery_stats['word_docs_recovered']}\n"
                              f"Content recovered: {recovery_stats['content_recovered']}\n"
                              f"Artifacts removed: {recovery_stats['artifacts_removed']}\n\n"
                              f"NOTE: Original files were deleted by virus")
        else:
            messagebox.showerror("Recovery Failed", "File recovery and cleaning process encountered errors.")

    def quarantine_threats(self):
        """Quarantine detected threats instead of recovering them"""
        if not self.detections:
            messagebox.showinfo("No Threats", "No threats detected to quarantine.")
            return
            
        result = messagebox.askyesno("Confirm Quarantine",
                                   f"Quarantine {len(self.detections)} detected threats?")
        if not result:
            return
            
        quarantined_count = 0
        
        for detection in self.detections:
            filepath = detection['file']
            if os.path.exists(filepath):
                reason = f"Threat: {detection['type']}"
                if 'document_type' in detection:
                    reason += f" ({detection['document_type']})"
                if detection.get('encrypted', False):
                    reason += " [ENCRYPTED FILE]"
                    
                if self.remover.quarantine_file(filepath, reason):
                    quarantined_count += 1

        self.log_to_console(f"Quarantined {quarantined_count} threats")
        self.update_quarantine_info()
        
        messagebox.showinfo("Quarantine Complete",
                          f"Successfully quarantined {quarantined_count} threats.")

    def update_quarantine_info(self):
        """Update quarantine information display"""
        self.quarantine_text.delete(1.0, tk.END)
        
        quarantine_dir = self.scanner.quarantine_dir
        
        if not os.path.exists(quarantine_dir):
            self.quarantine_text.insert(tk.END, "Quarantine directory not initialized.\n")
            return

        self.quarantine_text.insert(tk.END, f"Quarantine Directory: {quarantine_dir}\n")
        self.quarantine_text.insert(tk.END, "=" * 60 + "\n\n")
        
        try:
            files = os.listdir(quarantine_dir)
            quarantined_files = [f for f in files if not f.endswith('.metadata.json') and f != 'quarantine_info.json']
            
            self.quarantine_text.insert(tk.END, f"Total quarantined files: {len(quarantined_files)}\n\n")
            
            for filename in quarantined_files:
                filepath = os.path.join(quarantine_dir, filename)
                metadata_path = filepath + ".metadata.json"
                
                self.quarantine_text.insert(tk.END, f"File: {filename}\n")
                
                if os.path.exists(metadata_path):
                    try:
                        with open(metadata_path, 'r', encoding='utf-8') as f:
                            metadata = json.load(f)
                            
                        self.quarantine_text.insert(tk.END, f"  Original: {metadata.get('original_path', 'Unknown')}\n")
                        self.quarantine_text.insert(tk.END, f"  Quarantined: {metadata.get('quarantine_time', 'Unknown')}\n")
                        self.quarantine_text.insert(tk.END, f"  Reason: {metadata.get('reason', 'Unknown')}\n")
                        self.quarantine_text.insert(tk.END, f"  Size: {metadata.get('file_size', 'Unknown')} bytes\n")
                        
                        if metadata.get('encrypted_file', False):
                            self.quarantine_text.insert(tk.END, f"  Type: ENCRYPTED FILE\n")
                        
                    except Exception as e:
                        self.quarantine_text.insert(tk.END, f"  Error reading metadata: {str(e)}\n")
                        
                self.quarantine_text.insert(tk.END, "\n")
                
        except Exception as e:
            self.quarantine_text.insert(tk.END, f"Error reading quarantine directory: {str(e)}\n")

    def run(self):
        """Start the GUI"""
        print("DEBUG: Starting GUI main loop...")
        try:
            self.root.mainloop()
        except Exception as e:
            print(f"ERROR: GUI main loop error: {e}")
            raise

class CommandLineInterface:
    """Command-line interface for Odyssey Virus Hunter"""

    def __init__(self):
        self.scanner = AdvancedVirusScanner()
        self.remover = OdysseyVirusRemover(self.scanner)
        
    def display_banner(self):
        """Display application banner"""
        banner = """
ODYSSEY VIRUS HUNTER
===============================================================
    Encrypted File Detection & Recovery System
    National Teachers College - Information Assurance and Security 1 Finals Project
===============================================================

Target: Odyssey Virus
Encryption: ROT13 Decryption Engine
Features: Advanced Encrypted File Detection & Recovery
Behavior: Original files deleted - only encrypted files remain
Documents: Word, PDF, Excel, PowerPoint Support

"""
        print(banner)
        
    def run_interactive_scan(self):
        """Run interactive command-line scan"""
        self.display_banner()

        print("Initializing Antivirus Scanner...")
        print("⚠️  NOTICE: This version focuses on encrypted file recovery")
        print("⚠️  Original files were deleted by the virus")
        print(f"Document format support:")
        print(f"   Word (.docx): {'YES' if DOCX_AVAILABLE else 'NO'}")
        print(f"   PDF (.pdf): {'YES' if PDF_AVAILABLE else 'NO'}")
        print(f"   Excel (.xlsx): {'YES' if XLSX_AVAILABLE else 'NO'}")
        print(f"   PowerPoint (.pptx): {'YES' if PPTX_AVAILABLE else 'NO'}")
        
        # Initialize quarantine
        self.scanner.initialize_quarantine_system()
        
        while True:
            print("\n" + "="*60)
            print("Odyssey Virus Hunter - Select an option:")
            print("1. Comprehensive Encrypted & Infected File Scan")
            print("2. Quick Scan (Current Directory)")
            print("3. Recover & Clean All Files")
            print("4. View Quarantine")
            print("5. View Statistics")
            print("6. View Recovery Analysis")
            print("7. Exit")
            
            choice = input("\nEnter your choice (1-7): ").strip()
            
            if choice == '1':
                self.perform_comprehensive_scan()
            elif choice == '2':
                self.perform_quick_scan()
            elif choice == '3':
                self.perform_recovery()
            elif choice == '4':
                self.view_quarantine()
            elif choice == '5':
                self.view_statistics()
            elif choice == '6':
                self.view_recovery_analysis()
            elif choice == '7':
                print("\nThank you for using Odyssey Virus Hunter!")
                break
            else:
                print("Invalid choice. Please try again.")
                
    def perform_comprehensive_scan(self):
        """Perform comprehensive scan"""
        print("\nStarting comprehensive Odyssey virus scan (encrypted files, infected files, and artifacts)...")
        detections = self.scanner.perform_comprehensive_scan()
        self.display_scan_results(detections)
        
    def perform_quick_scan(self):
        """Perform quick scan"""
        print("\nStarting quick scan for encrypted and infected files...")
        detections = self.scanner.perform_comprehensive_scan(["."])
        self.display_scan_results(detections)
        
    def display_scan_results(self, detections):
        """Display scan results"""
        print(f"\nENCRYPTED FILE SCAN RESULTS")
        print("="*50)
        print(f"Threats detected: {len(detections)}")
        
        if detections:
            print("\nDETECTED THREATS:")

            # Group by type
            by_type = {}
            for detection in detections:
                det_type = detection['type']
                if det_type not in by_type:
                    by_type[det_type] = []
                by_type[det_type].append(detection)
            
            for det_type, type_detections in by_type.items():
                print(f"\n--- {det_type.upper().replace('_', ' ')} ({len(type_detections)}) ---")
                
                for i, detection in enumerate(type_detections, 1):
                    print(f"\n{i}. File: {detection['file']}")
                    print(f"   Confidence: {detection.get('confidence', 'MEDIUM')}")
                    
                    if 'document_type' in detection:
                        print(f"   Document Type: {detection['document_type']}")
                    
                    if 'virus_info' in detection:
                        info = detection['virus_info']
                        print(f"   Virus: {info['name']}")
                        print(f"   Risk: {info['risk_level']}")
                    
                    if detection.get('encrypted', False):
                        print(f"   Status: ENCRYPTED FILE (recoverable)")
                    
                    if 'virus_indicators' in detection and detection['virus_indicators']:
                        print(f"   Indicators: {len(detection['virus_indicators'])} found")
        else:
            print("\nNo encrypted files or threats detected. System appears clean!")

    def perform_recovery(self):
        """Perform file recovery and cleaning"""
        print("\nStarting file recovery and cleaning process...")
        
        target_dir = "odyssey_test_environment"
        if not os.path.exists(target_dir):
            print(f"Target directory '{target_dir}' not found.")
            return
            
        success = self.remover.perform_complete_removal(target_dir)
        
        if success:
            print("\nFile recovery and cleaning completed successfully!")
            stats = self.remover.recovery_statistics
            print(f"\nRecovery & Cleaning Statistics:")
            print(f"   Infected Files Processed: {stats['infected_files_processed']}")
            print(f"   Files Cleaned: {stats['files_cleaned']}")
            print(f"   Encrypted Files Processed: {stats['encrypted_files_processed']}")
            print(f"   Files Decrypted: {stats['files_decrypted']}")
            print(f"   Documents Recovered: {stats['documents_recovered']}")
            print(f"   Word Documents Recovered: {stats['word_docs_recovered']}")
            print(f"   Content Recovered: {stats['content_recovered']}")
            print(f"   Artifacts Removed: {stats['artifacts_removed']}")
            print("\nNOTE: Original files were deleted by virus - recovered content and cleaned files available")
        else:
            print("\nFile recovery and cleaning encountered errors.")

    def view_quarantine(self):
        """View quarantine information"""
        print(f"\nQUARANTINE INFORMATION")
        print("="*50)
        
        quarantine_dir = self.scanner.quarantine_dir
        
        if not os.path.exists(quarantine_dir):
            print("Quarantine directory not initialized.")
            return
            
        try:
            files = os.listdir(quarantine_dir)
            quarantined_files = [f for f in files if not f.endswith('.metadata.json') and f != 'quarantine_info.json']
            
            print(f"Directory: {quarantine_dir}")
            print(f"Quarantined files: {len(quarantined_files)}")
            
            if quarantined_files:
                print("\nQuarantined Files:")
                for filename in quarantined_files[:10]:  # Show first 10
                    print(f"  {filename}")
                    
                if len(quarantined_files) > 10:
                    print(f"  ... and {len(quarantined_files) - 10} more files")
                    
        except Exception as e:
            print(f"Error reading quarantine: {str(e)}")

    def view_statistics(self):
        """View scan statistics"""
        print(f"\nANTIVIRUS STATISTICS")
        print("="*50)
        
        stats = self.scanner.scan_statistics
        recovery_stats = self.remover.recovery_statistics

        print(f"Scan Statistics:")
        print(f"  Files Scanned: {stats['files_scanned']}")
        print(f"  Encrypted Files Found: {stats['encrypted_files_found']}")
        print(f"  Infected Files Found: {stats['infected_files_found']}")
        print(f"  Artifacts Found: {stats['artifacts_found']}")
        print(f"  Threats Detected: {stats['threats_detected']}")
        print(f"  Last Scan Duration: {stats['scan_duration']:.2f} seconds")
        
        print(f"\nDocument Types Scanned:")
        if stats['document_types_scanned']:
            for doc_type, count in stats['document_types_scanned'].items():
                print(f"  {doc_type}: {count}")
        else:
            print("  No documents processed yet")

        print(f"\nRecovery Statistics:")
        print(f"  Infected Files Processed: {recovery_stats['infected_files_processed']}")
        print(f"  Files Cleaned: {recovery_stats['files_cleaned']}")
        print(f"  Encrypted Files Processed: {recovery_stats['encrypted_files_processed']}")
        print(f"  Files Decrypted: {recovery_stats['files_decrypted']}")
        print(f"  Documents Recovered: {recovery_stats['documents_recovered']}")
        print(f"  Word Documents Recovered: {recovery_stats['word_docs_recovered']}")
        print(f"  Content Recovered: {recovery_stats['content_recovered']}")
        print(f"  Artifacts Removed: {recovery_stats['artifacts_removed']}")
        print(f"  Files Quarantined: {recovery_stats['files_quarantined']}")
        
    def view_recovery_analysis(self):
        """View encrypted file recovery analysis information"""
        print(f"\nENCRYPTED FILE RECOVERY ANALYSIS")
        print("="*50)
        
        print("Library Availability:")
        print(f"  Word Document Support: {'Available' if DOCX_AVAILABLE else 'Not Available'}")
        print(f"  PDF Support: {'Available' if PDF_AVAILABLE else 'Not Available'}")
        print(f"  Excel Support: {'Available' if XLSX_AVAILABLE else 'Not Available'}")
        print(f"  PowerPoint Support: {'Available' if PPTX_AVAILABLE else 'Not Available'}")
        
        stats = self.scanner.scan_statistics
        
        print(f"\nFile Processing Statistics:")
        print(f"  Total Files Scanned: {stats['files_scanned']}")
        print(f"  Encrypted Files Found: {stats['encrypted_files_found']}")
        print(f"  Infected Files Found: {stats['infected_files_found']}")
        print(f"  Artifacts Found: {stats['artifacts_found']}")
        print(f"  Detection Rate: {(stats['threats_detected'] / max(stats['files_scanned'], 1) * 100):.1f}%")
        print(f"  Encrypted File Rate: {(stats['encrypted_files_found'] / max(stats['files_scanned'], 1) * 100):.1f}%")
        print(f"  Infected File Rate: {(stats['infected_files_found'] / max(stats['files_scanned'], 1) * 100):.1f}%")
        
        if stats['document_types_scanned']:
            print(f"\nDocument Types Processed:")
            for doc_type, count in stats['document_types_scanned'].items():
                print(f"  • {doc_type}: {count} files")
        
        recovery_stats = self.remover.recovery_statistics
        if recovery_stats['encrypted_files_processed'] > 0 or recovery_stats['infected_files_processed'] > 0:
            print(f"\nRecovery Analysis:")
            print(f"  Infected Files Processed: {recovery_stats['infected_files_processed']}")
            print(f"  Files Cleaned: {recovery_stats['files_cleaned']}")
            print(f"  Encrypted Files Processed: {recovery_stats['encrypted_files_processed']}")
            print(f"  Decryption Success Rate: {(recovery_stats['content_recovered'] / max(recovery_stats['encrypted_files_processed'], 1) * 100):.1f}%")
            print(f"  Cleaning Success Rate: {(recovery_stats['files_cleaned'] / max(recovery_stats['infected_files_processed'], 1) * 100):.1f}%")
            print(f"  Word Document Recovery: {recovery_stats['word_docs_recovered']} files")
            print(f"  Text File Recovery: {recovery_stats['files_decrypted']} files")
        
        print(f"\nRecovery Notes:")
        print("  • Original files were deleted by the virus")
        print("  • Encrypted files can be decrypted using ROT13")
        print("  • Infected files can be cleaned by removing virus signatures")
        print("  • Backup files may be available for infected file restoration")
        print("  • ROT13 decryption restores original content")
        print("  • Recovered and cleaned files are marked with metadata")

def main():
    """Main entry point for Odyssey Virus Hunter"""
    print("DEBUG: Starting main function ...")
    print("Odyssey Virus Hunter - Initializing...")
    
    # Display library status
    print("\nDocument Processing Library Status:")
    print(f"   python-docx (Word): {'Available' if DOCX_AVAILABLE else 'Not Available - install with: pip install python-docx'}")
    print(f"   PyPDF2 (PDF): {'Available' if PDF_AVAILABLE else 'Not Available - install with: pip install PyPDF2'}")
    print(f"   openpyxl (Excel): {'Available' if XLSX_AVAILABLE else 'Not Available - install with: pip install openpyxl'}")
    print(f"   python-pptx (PowerPoint): {'Available' if PPTX_AVAILABLE else 'Not Available - install with: pip install python-pptx'}")
    
    missing_libs = []
    if not DOCX_AVAILABLE:
        missing_libs.append("python-docx")
    if not PDF_AVAILABLE:
        missing_libs.append("PyPDF2")
    if not XLSX_AVAILABLE:
        missing_libs.append("openpyxl")
    if not PPTX_AVAILABLE:
        missing_libs.append("python-pptx")
    
    if missing_libs:
        print(f"\nOptional libraries missing: {', '.join(missing_libs)}")
        print("Install all with: pip install python-docx PyPDF2 openpyxl python-pptx")
        print("(Antivirus will work with reduced functionality)\n")
    else:
        print("\nAll document processing libraries available!\n")
    
    print("🎯 VERSION FEATURES:")
    print("   • Focus on encrypted file detection and recovery")
    print("   • Original files were deleted by virus")
    print("   • ROT13 decryption for content recovery")
    print("   • Multi-format document processing")
    print("   • Enhanced artifact detection")
    print()
    
    # Check for GUI availability with improved error handling
    if GUI_AVAILABLE:
        print("GUI available. Choose interface:")
        print("1. Graphical Interface (Recommended)")
        print("2. Command Line Interface")
        
        max_attempts = 3
        attempt = 0
        
        while attempt < max_attempts:
            try:
                choice = input("\nEnter choice (1 or 2): ").strip()
                
                if choice == '1':
                    print("Starting GUI mode...")
                    print("DEBUG: About to create AntivirusGUI instance...")
                    
                    try:
                        antivirus = AntivirusGUI()
                        print("DEBUG: AntivirusGUI created successfully")
                        antivirus.run()
                        print("DEBUG: GUI run completed")
                    except Exception as e:
                        print(f"GUI Error: {str(e)}")
                        print("DEBUG: Full GUI error details:")
                        import traceback
                        traceback.print_exc()
                        print("Falling back to CLI mode...")
                        cli = CommandLineInterface()
                        cli.run_interactive_scan()
                    break
                elif choice == '2':
                    print("Starting CLI mode...")
                    cli = CommandLineInterface()
                    cli.run_interactive_scan()
                    break
                else:
                    print("Invalid choice. Please enter 1 or 2.")
                    attempt += 1
                    
            except KeyboardInterrupt:
                print("\nExiting...")
                break
            except Exception as e:
                print(f"Input error: {e}")
                attempt += 1
                
        if attempt >= max_attempts:
            print("Too many invalid attempts. Starting CLI mode...")
            cli = CommandLineInterface()
            cli.run_interactive_scan()
                
    else:
        # GUI not available, use CLI
        print("GUI not available, using command line interface...")
        cli = CommandLineInterface()
        cli.run_interactive_scan()

if __name__ == "__main__":
    print("DEBUG: Script starting...")
    try:
        main()
    except Exception as e:
        print(f"FATAL ERROR: {str(e)}")
        print("DEBUG: Full error details:")
        import traceback
        traceback.print_exc()
        input("Press Enter to exit...")
    finally:
        print("DEBUG: Script ending...")

def main():
    """Main entry point for Odyssey Virus Hunter"""
    print("DEBUG: Starting main function...")
    print("Odyssey Virus Hunter - Initializing...")
    
    # Display library status
    print("\nDocument Processing Library Status:")
    print(f"   python-docx (Word): {'Available' if DOCX_AVAILABLE else 'Not Available - install with: pip install python-docx'}")
    print(f"   PyPDF2 (PDF): {'Available' if PDF_AVAILABLE else 'Not Available - install with: pip install PyPDF2'}")
    print(f"   openpyxl (Excel): {'Available' if XLSX_AVAILABLE else 'Not Available - install with: pip install openpyxl'}")
    print(f"   python-pptx (PowerPoint): {'Available' if PPTX_AVAILABLE else 'Not Available - install with: pip install python-pptx'}")
    
    missing_libs = []
    if not DOCX_AVAILABLE:
        missing_libs.append("python-docx")
    if not PDF_AVAILABLE:
        missing_libs.append("PyPDF2")
    if not XLSX_AVAILABLE:
        missing_libs.append("openpyxl")
    if not PPTX_AVAILABLE:
        missing_libs.append("python-pptx")
    
    if missing_libs:
        print(f"\nOptional libraries missing: {', '.join(missing_libs)}")
        print("Install all with: pip install python-docx PyPDF2 openpyxl python-pptx")
        print("(Antivirus will work with reduced functionality)\n")
    else:
        print("\nAll document processing libraries available!\n")

    print("🎯 VERSION FEATURES:")
    print("   • Focus on encrypted file detection and recovery")
    print("   • Original files were deleted by virus")
    print("   • ROT13 decryption for content recovery")
    print("   • Multi-format document processing")
    print("   • Enhanced artifact detection")
    print()
    
    # Check for GUI availability with improved error handling
    if GUI_AVAILABLE:
        print("GUI available. Choose interface:")
        print("1. Graphical Interface (Recommended)")
        print("2. Command Line Interface")
        
        max_attempts = 3
        attempt = 0
        
        while attempt < max_attempts:
            try:
                choice = input("\nEnter choice (1 or 2): ").strip()
                
                if choice == '1':
                    print("Starting GUI mode...")
                    print("DEBUG: About to create AntivirusGUI instance...")
                    
                    try:
                        antivirus = AntivirusGUI()
                        print("DEBUG: AntivirusGUI created successfully")
                        antivirus.run()
                        print("DEBUG: GUI run completed")
                    except Exception as e:
                        print(f"GUI Error: {str(e)}")
                        print("DEBUG: Full GUI error details:")
                        import traceback
                        traceback.print_exc()
                        print("Falling back to CLI mode...")
                        cli = CommandLineInterface()
                        cli.run_interactive_scan()
                    break
                elif choice == '2':
                    print("Starting CLI mode...")
                    cli = CommandLineInterface()
                    cli.run_interactive_scan()
                    break
                else:
                    print("Invalid choice. Please enter 1 or 2.")
                    attempt += 1
                    
            except KeyboardInterrupt:
                print("\nExiting...")
                break
            except Exception as e:
                print(f"Input error: {e}")
                attempt += 1
                
        if attempt >= max_attempts:
            print("Too many invalid attempts. Starting CLI mode...")
            cli = CommandLineInterface()
            cli.run_interactive_scan()
                
    else:
        # GUI not available, use CLI
        print("GUI not available, using command line interface...")
        cli = CommandLineInterface()
        cli.run_interactive_scan()

if __name__ == "__main__":
    print("DEBUG: Script starting...")
    try:
        main()
    except Exception as e:
        print(f"FATAL ERROR: {str(e)}")
        print("DEBUG: Full error details:")
        import traceback
        traceback.print_exc()
        input("Press Enter to exit...")
    finally:
        print("DEBUG: Script ending...")